from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("类继承关系图")
run.font.size = Pt(16)
run.font.bold = True

doc.add_paragraph()  # 空行

# ── 用表格画图 ──
table = doc.add_table(rows=1, cols=5)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 样式
def cell(text, bg=None, bold=False, size=10):
    c = table.cell(0, 0)  # placeholder, we'll use merge
    return text, bg, bold, size

# 重新来：直接用段落 + 等宽字体，简单可靠
doc.add_paragraph()  # clear

# 图例区域
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = legend.add_run("── 实线 = 继承(extends)　　　- - 虚线 = 组合(has-a)")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 画图的函数
def add_diagram_line(text, font_name='Courier New', size=10, bold=False, color=None, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# ===== 类图 =====
lines = [
    ("                      ┌──────────────────────┐", None),
    ("                      │   Entity  (抽象基类)    │", (0x15, 0x65, 0xC0)),
    ("                      │   坐标 / 符号 / 存活    │", None),
    ("                      │   update()  纯虚函数    │", None),
    ("                      │   draw()    虚函数      │", None),
    ("                      └──┬────────┬────────┬──┘", None),
    ("           extends  ──┐   │        │   ┌── extends", None),
    ("                      │   │        │   │", None),
    ("         ┌────────────┘   │        └───┼────────────┐", None),
    ("         ▼                ▼            ▼            ▼", None),
    ("   ┌──────────┐    ┌──────────┐   ┌──────────┐", None),
    ("   │  Player  │    │  Enemy   │   │  Bullet  │", (0x2E, 0x7D, 0x32)),
    ("   │ 移动/射击 │    │  向下掉落  │   │  向上飞行  │", None),
    ("   │ 显示: A  │    │ 显示: V  │   │ 显示: |  │", (0xE6, 0x51, 0x00)),
    ("   └────┬─────┘    └────┬─────┘   └─────┬─────┘", None),
    ("        │               │               │", None),
    ("        │     has-a     │     has-a     │", None),
    ("        └───────────────┼───────────────┘", None),
    ("                        ▼", None),
    ("              ┌──────────────────┐", None),
    ("              │      Game        │", (0x6A, 0x1B, 0x9A)),
    ("              │     游戏主控      │", None),
    ("              │  主循环/碰撞/渲染  │", None),
    ("              └──────────────────┘", None),
]

for text, color in lines:
    add_diagram_line(text, color=color)

doc.add_paragraph()  # 空行

# 类职责说明表
doc.add_heading('各类职责说明', level=2)

table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

headers = ['类名', '继承自', '职责', '关键方法']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

data = [
    ['Entity', '—（基类）', '所有实体的公共抽象', 'update(), draw(), collidesWith()'],
    ['Player', 'Entity', '玩家控制的战机，支持移动和射击', 'move(), tryFire()'],
    ['Enemy',  'Entity', '敌机，向下移动，出界即死', 'update()'],
    ['Bullet', 'Entity', '子弹，向上移动，出界即死', 'update()'],
    ['Game',   '—（独立）', '游戏主控，持有所有实体并驱动循环', 'run(), render(), checkCollisions()'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, text in enumerate(row_data):
        cell = table.cell(row_idx + 1, col_idx)
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.save('类继承关系图.docx')
print("Saved: 类继承关系图.docx")
